"""
Document Processor - 실제 문서 처리 시스템

문서 로딩, 전처리, 메타데이터 추출을 담당하는 
실제 동작하는 문서 처리 엔진.

Features:
- 다양한 문서 형식 지원 (PDF, DOCX, TXT, MD)
- 한국어 텍스트 정규화
- 메타데이터 자동 추출
- 뿌리산업 특화 전처리
- 이미지/표 추출 지원
"""

import asyncio
import logging
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import hashlib
import mimetypes
import json

# 문서 처리 라이브러리
try:
    import PyPDF2
    import docx
    from PIL import Image
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    EMBEDDING_AVAILABLE = True
except ImportError:
    EMBEDDING_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """처리된 문서"""
    id: str
    title: str
    content: str
    original_content: str
    file_type: str
    file_size: int
    language: str
    namespace: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    extracted_entities: List[str] = field(default_factory=list)
    industry_terms: List[str] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    created_at: datetime = field(default_factory=datetime.now)
    processed_at: datetime = field(default_factory=datetime.now)


@dataclass
class ProcessingConfig:
    """문서 처리 설정"""
    extract_images: bool = True
    extract_tables: bool = True
    normalize_korean: bool = True
    remove_noise: bool = True
    min_text_length: int = 10
    max_file_size_mb: int = 50
    supported_formats: List[str] = field(default_factory=lambda: [
        '.pdf', '.docx', '.txt', '.md', '.html', '.rtf'
    ])


class DocumentProcessor:
    """
    실제 문서 처리 엔진
    
    다양한 형식의 문서를 로딩하고 전처리하여
    RAG 시스템에서 사용할 수 있는 형태로 변환.
    """
    
    def __init__(
        self,
        config_manager,
        korean_optimizer=None,
        config: Optional[ProcessingConfig] = None
    ):
        self.config_manager = config_manager
        self.korean_optimizer = korean_optimizer
        self.config = config or ProcessingConfig()
        
        # 임베딩 모델 초기화
        self.embedding_model = None
        if EMBEDDING_AVAILABLE:
            try:
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                logger.info("임베딩 모델 로드 완료: all-MiniLM-L6-v2")
            except Exception as e:
                logger.warning(f"임베딩 모델 로드 실패: {e}")
        
        # 뿌리산업 키워드 매핑
        self.industry_keywords = self._initialize_industry_keywords()
        
        # 처리 통계
        self.processing_stats = {
            "total_processed": 0,
            "success_count": 0,
            "error_count": 0,
            "avg_processing_time": 0.0,
            "total_size_mb": 0.0
        }
        
        logger.info("Document Processor 초기화 완료")
    
    def _initialize_industry_keywords(self) -> Dict[str, List[str]]:
        """뿌리산업 키워드 초기화"""
        
        return {
            "주조": [
                "주조", "캐스팅", "용탕", "응고", "주형", "모래주형", "금속주형",
                "다이캐스팅", "원심주조", "정밀주조", "로스트왁스", "셸몰드"
            ],
            "금형": [
                "금형", "다이", "몰드", "프레스", "성형", "사출", "압출",
                "블로우", "스탬핑", "드로잉", "벤딩", "펀칭"
            ],
            "소성가공": [
                "소성가공", "단조", "압연", "인발", "전조", "압출", "성형",
                "냉간가공", "열간가공", "자유단조", "형단조", "정밀단조"
            ],
            "용접": [
                "용접", "접합", "아크용접", "가스용접", "저항용접", "레이저용접",
                "전자빔용접", "마찰용접", "브레이징", "솔더링", "용접봉"
            ],
            "표면처리": [
                "표면처리", "도금", "코팅", "양극산화", "침탄", "질화",
                "쇼트피닝", "연마", "화성처리", "PVD", "CVD", "전해연마"
            ],
            "열처리": [
                "열처리", "소입", "소성", "담금질", "풀림", "노멀라이징",
                "템퍼링", "침탄", "질화", "고주파", "진공열처리"
            ]
        }
    
    async def process_file(
        self,
        file_path: Union[str, Path],
        namespace: str = "default",
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> Optional[ProcessedDocument]:
        """
        파일 처리
        
        Args:
            file_path: 처리할 파일 경로
            namespace: 문서가 속할 네임스페이스
            custom_metadata: 추가 메타데이터
            
        Returns:
            ProcessedDocument: 처리된 문서 객체
        """
        
        start_time = asyncio.get_event_loop().time()
        
        try:
            file_path = Path(file_path)
            
            # 파일 존재 확인
            if not file_path.exists():
                logger.error(f"파일을 찾을 수 없음: {file_path}")
                return None
            
            # 파일 크기 확인
            file_size = file_path.stat().st_size
            if file_size > self.config.max_file_size_mb * 1024 * 1024:
                logger.error(f"파일 크기 초과: {file_size / 1024 / 1024:.1f}MB")
                return None
            
            # 파일 형식 확인
            file_extension = file_path.suffix.lower()
            if file_extension not in self.config.supported_formats:
                logger.error(f"지원하지 않는 파일 형식: {file_extension}")
                return None
            
            # 파일 내용 추출
            content = await self._extract_content(file_path)
            if not content or len(content) < self.config.min_text_length:
                logger.warning(f"추출된 텍스트가 너무 짧음: {len(content) if content else 0}자")
                return None
            
            # 문서 ID 생성
            doc_id = self._generate_document_id(file_path, content)
            
            # 메타데이터 추출
            metadata = await self._extract_metadata(file_path, content)
            if custom_metadata:
                metadata.update(custom_metadata)
            
            # 한국어 처리
            processed_content = content
            industry_terms = []
            
            if self.korean_optimizer and self.config.normalize_korean:
                korean_result = await self.korean_optimizer.process_korean_text(content)
                processed_content = korean_result.normalized_text
                industry_terms = korean_result.industry_terms
            
            # 엔티티 추출
            entities = await self._extract_entities(processed_content)
            
            # 이미지/표 추출
            images = []
            tables = []
            
            if file_extension == '.pdf' and PDF_AVAILABLE:
                if self.config.extract_images:
                    images = await self._extract_images_from_pdf(file_path)
                if self.config.extract_tables:
                    tables = await self._extract_tables_from_pdf(file_path)
            
            # 처리된 문서 객체 생성
            processed_doc = ProcessedDocument(
                id=doc_id,
                title=metadata.get("title", file_path.stem),
                content=processed_content,
                original_content=content,
                file_type=file_extension,
                file_size=file_size,
                language=metadata.get("language", "ko"),
                namespace=namespace,
                metadata=metadata,
                extracted_entities=entities,
                industry_terms=industry_terms,
                images=images,
                tables=tables
            )
            
            # 처리 통계 업데이트
            processing_time = asyncio.get_event_loop().time() - start_time
            self._update_processing_stats(file_size, processing_time, True)
            
            logger.info(f"문서 처리 완료: {file_path.name} ({processing_time:.2f}초)")
            return processed_doc
            
        except Exception as e:
            processing_time = asyncio.get_event_loop().time() - start_time
            self._update_processing_stats(0, processing_time, False)
            
            logger.error(f"문서 처리 실패 ({file_path}): {e}")
            return None
    
    async def process_text(
        self,
        text: str,
        title: str = "텍스트 문서",
        namespace: str = "default",
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> Optional[ProcessedDocument]:
        """
        텍스트 직접 처리
        
        Args:
            text: 처리할 텍스트
            title: 문서 제목
            namespace: 네임스페이스
            custom_metadata: 추가 메타데이터
            
        Returns:
            ProcessedDocument: 처리된 문서 객체
        """
        
        start_time = asyncio.get_event_loop().time()
        
        try:
            if len(text) < self.config.min_text_length:
                logger.warning(f"텍스트가 너무 짧음: {len(text)}자")
                return None
            
            # 문서 ID 생성
            doc_id = hashlib.md5(text.encode('utf-8')).hexdigest()
            
            # 메타데이터 구성
            metadata = {
                "title": title,
                "source": "direct_text",
                "language": "ko",
                "created_at": datetime.now().isoformat(),
                **(custom_metadata or {})
            }
            
            # 한국어 처리
            processed_content = text
            industry_terms = []
            
            if self.korean_optimizer and self.config.normalize_korean:
                korean_result = await self.korean_optimizer.process_korean_text(text)
                processed_content = korean_result.normalized_text
                industry_terms = korean_result.industry_terms
            
            # 엔티티 추출
            entities = await self._extract_entities(processed_content)
            
            # 처리된 문서 객체 생성
            processed_doc = ProcessedDocument(
                id=doc_id,
                title=title,
                content=processed_content,
                original_content=text,
                file_type="text",
                file_size=len(text.encode('utf-8')),
                language="ko",
                namespace=namespace,
                metadata=metadata,
                extracted_entities=entities,
                industry_terms=industry_terms
            )
            
            # 처리 통계 업데이트
            processing_time = asyncio.get_event_loop().time() - start_time
            self._update_processing_stats(len(text.encode('utf-8')), processing_time, True)
            
            logger.info(f"텍스트 처리 완료: {title} ({processing_time:.2f}초)")
            return processed_doc
            
        except Exception as e:
            processing_time = asyncio.get_event_loop().time() - start_time
            self._update_processing_stats(0, processing_time, False)
            
            logger.error(f"텍스트 처리 실패: {e}")
            return None
    
    async def _extract_content(self, file_path: Path) -> Optional[str]:
        """파일에서 텍스트 추출"""
        
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.txt':
                return await self._extract_from_txt(file_path)
            elif file_extension == '.md':
                return await self._extract_from_markdown(file_path)
            elif file_extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return await self._extract_from_docx(file_path)
            elif file_extension == '.html':
                return await self._extract_from_html(file_path)
            else:
                # 기본적으로 텍스트 파일로 시도
                return await self._extract_from_txt(file_path)
                
        except Exception as e:
            logger.error(f"내용 추출 실패 ({file_path}): {e}")
            return None
    
    async def _extract_from_txt(self, file_path: Path) -> str:
        """TXT 파일에서 텍스트 추출"""
        
        encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"지원하는 인코딩으로 파일을 읽을 수 없음: {file_path}")
    
    async def _extract_from_markdown(self, file_path: Path) -> str:
        """Markdown 파일에서 텍스트 추출"""
        
        content = await self._extract_from_txt(file_path)
        
        # 마크다운 문법 제거 (간단한 버전)
        import re
        
        # 헤더 제거
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        
        # 링크 제거
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        
        # 강조 제거
        content = re.sub(r'\*\*([^\*]+)\*\*', r'\1', content)
        content = re.sub(r'\*([^\*]+)\*', r'\1', content)
        
        # 코드 블록 제거
        content = re.sub(r'```[^`]*```', '', content, flags=re.DOTALL)
        content = re.sub(r'`([^`]+)`', r'\1', content)
        
        return content.strip()
    
    async def _extract_from_pdf(self, file_path: Path) -> str:
        """PDF 파일에서 텍스트 추출"""
        
        if not PDF_AVAILABLE:
            raise ImportError("PDF 처리 라이브러리가 설치되지 않음 (PyPDF2, PyMuPDF)")
        
        text_content = []
        
        try:
            # PyMuPDF 사용 (더 정확한 텍스트 추출)
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content.append(page.get_text())
            
            doc.close()
            
        except Exception:
            # 대안으로 PyPDF2 사용
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
        
        return '\n'.join(text_content).strip()
    
    async def _extract_from_docx(self, file_path: Path) -> str:
        """DOCX 파일에서 텍스트 추출"""
        
        try:
            doc = docx.Document(file_path)
            
            text_content = []
            
            # 단락 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 표 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content).strip()
            
        except ImportError:
            raise ImportError("python-docx 라이브러리가 설치되지 않음")
    
    async def _extract_from_html(self, file_path: Path) -> str:
        """HTML 파일에서 텍스트 추출"""
        
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # 스크립트/스타일 태그 제거
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 텍스트 추출
            text = soup.get_text()
            
            # 연속 공백 정리
            import re
            text = re.sub(r'\s+', ' ', text)
            
            return text.strip()
            
        except ImportError:
            raise ImportError("beautifulsoup4 라이브러리가 설치되지 않음")
    
    def _generate_document_id(self, file_path: Path, content: str) -> str:
        """문서 ID 생성"""
        
        # 파일 경로 + 내용 해시 조합
        file_info = f"{file_path.name}:{file_path.stat().st_mtime}"
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()[:8]
        
        return f"{hashlib.md5(file_info.encode('utf-8')).hexdigest()[:8]}_{content_hash}"
    
    async def _extract_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """메타데이터 추출"""
        
        metadata = {
            "file_name": file_path.name,
            "file_path": str(file_path.absolute()),
            "file_extension": file_path.suffix,
            "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
            "mime_type": mimetypes.guess_type(str(file_path))[0],
            "character_count": len(content),
            "word_count": len(content.split()),
            "line_count": content.count('\n') + 1,
        }
        
        # 언어 감지 (간단한 버전)
        korean_chars = len([c for c in content if '\uac00' <= c <= '\ud7af'])
        if korean_chars > len(content) * 0.1:  # 10% 이상 한글
            metadata["language"] = "ko"
        else:
            metadata["language"] = "en"
        
        # 뿌리산업 도메인 감지
        detected_domains = []
        for domain, keywords in self.industry_keywords.items():
            for keyword in keywords:
                if keyword in content:
                    detected_domains.append(domain)
                    break
        
        if detected_domains:
            metadata["detected_domains"] = detected_domains
            metadata["primary_domain"] = detected_domains[0]
        
        return metadata
    
    async def _extract_entities(self, content: str) -> List[str]:
        """엔티티 추출 (간단한 버전)"""
        
        entities = []
        
        # 뿌리산업 용어 추출
        for domain, keywords in self.industry_keywords.items():
            for keyword in keywords:
                if keyword in content:
                    entities.append(keyword)
        
        # 숫자 패턴 추출 (온도, 압력 등)
        import re
        
        # 온도 패턴
        temp_patterns = re.findall(r'\d+(?:\.\d+)?°C|\d+(?:\.\d+)?도', content)
        entities.extend([f"온도:{temp}" for temp in temp_patterns[:5]])
        
        # 압력 패턴  
        pressure_patterns = re.findall(r'\d+(?:\.\d+)?(?:MPa|kPa|bar|atm)', content)
        entities.extend([f"압력:{pressure}" for pressure in pressure_patterns[:5]])
        
        # 중복 제거
        return list(set(entities))
    
    async def _extract_images_from_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """PDF에서 이미지 추출"""
        
        if not PDF_AVAILABLE:
            return []
        
        images = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    # 이미지 메타데이터만 저장 (실제 이미지 데이터는 제외)
                    images.append({
                        "page": page_num + 1,
                        "index": img_index,
                        "width": img[2] if len(img) > 2 else 0,
                        "height": img[3] if len(img) > 3 else 0,
                        "type": "embedded_image"
                    })
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PDF 이미지 추출 실패: {e}")
        
        return images[:10]  # 최대 10개
    
    async def _extract_tables_from_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """PDF에서 표 추출"""
        
        if not PDF_AVAILABLE:
            return []
        
        tables = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 표 감지 (간단한 버전 - 실제로는 더 정교한 라이브러리 필요)
                text = page.get_text()
                
                # 탭이나 공백으로 구분된 데이터 감지
                lines = text.split('\n')
                potential_tables = []
                
                for line in lines:
                    if '\t' in line or len(line.split()) > 3:
                        potential_tables.append(line)
                
                if len(potential_tables) > 2:  # 최소 3줄 이상
                    tables.append({
                        "page": page_num + 1,
                        "row_count": len(potential_tables),
                        "preview": potential_tables[:3],  # 처음 3줄만
                        "type": "detected_table"
                    })
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PDF 표 추출 실패: {e}")
        
        return tables[:5]  # 최대 5개
    
    def _update_processing_stats(self, file_size: int, processing_time: float, success: bool):
        """처리 통계 업데이트"""
        
        self.processing_stats["total_processed"] += 1
        
        if success:
            self.processing_stats["success_count"] += 1
            self.processing_stats["total_size_mb"] += file_size / (1024 * 1024)
        else:
            self.processing_stats["error_count"] += 1
        
        # 평균 처리 시간 업데이트
        total = self.processing_stats["total_processed"]
        current_avg = self.processing_stats["avg_processing_time"]
        self.processing_stats["avg_processing_time"] = \
            (current_avg * (total - 1) + processing_time) / total
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """처리 통계 조회"""
        
        return {
            **self.processing_stats,
            "success_rate": (
                self.processing_stats["success_count"] / 
                max(1, self.processing_stats["total_processed"])
            ),
            "avg_file_size_mb": (
                self.processing_stats["total_size_mb"] / 
                max(1, self.processing_stats["success_count"])
            ),
            "supported_formats": self.config.supported_formats,
            "embedding_model_available": self.embedding_model is not None,
            "pdf_processing_available": PDF_AVAILABLE
        }
    
    async def cleanup(self):
        """Document Processor 정리"""
        
        # 임베딩 모델 정리
        if self.embedding_model:
            del self.embedding_model
            self.embedding_model = None
        
        logger.info("Document Processor 정리 완료")